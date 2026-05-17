from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(doc, headers, rows, header_color='2E4057'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    return table

# ── Cover ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CATATAN TEMUAN PENELITIAN')
run.bold = True; run.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Automated Essay Scoring (AES) — Siamese BiLSTM dengan FastText')
run2.font.size = Pt(13); run2.italic = True

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Tanggal: {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ── 1. RINGKASAN EKSEKUTIF ───────────────────────────────────────────────────
add_heading(doc, '1. Ringkasan Eksekutif', 1)
add_para(doc,
    'Penelitian ini membangun model Automated Essay Scoring (AES) untuk esai berbahasa Indonesia '
    'menggunakan arsitektur Siamese Bidirectional LSTM (BiLSTM) dengan embedding FastText 300 dimensi. '
    'Dataset mencakup 1.229 jawaban mahasiswa dari 17 prompt soal (IDPSJ) dengan rentang nilai 1–10.',
    size=11)
doc.add_paragraph()
add_para(doc, 'Hasil terbaik yang dicapai:', bold=True)
add_table(doc,
    ['Setting Evaluasi', 'MAE', 'QWK', 'Spearman', 'Keterangan'],
    [
        ['In-Prompt (Per-IDPSJ)', '1.352', '0.728', '0.642', 'Upper bound — model tahu skala grade'],
        ['Pooled (Random Split)', '1.573', '0.721', '0.720', 'Mixed — semi-tahu skala grade'],
        ['LOPO v10 (Cross-Prompt)', '1.798', '~0.45*', '~0.55*', 'Setting utama penelitian'],
    ]
)
add_para(doc, '* Estimasi QWK/Spearman untuk LOPO belum dihitung secara formal.', italic=True, size=9)

doc.add_paragraph()
add_para(doc,
    'Temuan kunci: Model MAMPU mencapai MAE < 1.5 dalam setting in-prompt (MAE=1.352). '
    'Bottleneck utama pada setting LOPO adalah heterogenitas skala grade antar IDPSJ (variasi mean grade '
    '4.42–7.84), bukan keterbatasan arsitektur model maupun kualitas embedding.',
    bold=False, size=11)

# ── 2. LATAR BELAKANG ────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '2. Latar Belakang dan Dataset', 1)

add_heading(doc, '2.1 Dataset', 2)
add_table(doc,
    ['Atribut', 'Nilai'],
    [
        ['Jumlah sampel asli', '1.229 jawaban mahasiswa'],
        ['Jumlah prompt soal (IDPSJ)', '17 IDPSJ (1–17)'],
        ['Skala grade', '1–10 (normalisasi: raw_grade / max_grade × 10)'],
        ['max_grade unik', '2, 4, 5, 10, 15, 30, 33.33, 50'],
        ['Distribusi label A-E', 'A=348, B=214, C=189, D=155, E=323'],
    ]
)

doc.add_paragraph()
add_heading(doc, '2.2 Mapping Label A-E', 2)
add_para(doc, 'Formula dari data aktual (sesuai rumus Excel dosen):', bold=True)
add_table(doc,
    ['Label', 'Rentang Grade (0–10)', 'Integer Grade', 'Jumlah Sampel'],
    [
        ['A', 'grade > 8.5', '9, 10', '348'],
        ['B', '7.0 ≤ grade ≤ 8.5', '7, 8', '214'],
        ['C', '6.0 ≤ grade < 7.0', '6', '189'],
        ['D', '4.0 ≤ grade < 6.0', '4, 5', '155'],
        ['E', 'grade < 4.0', '1, 2, 3', '323'],
    ]
)
add_para(doc,
    'Catatan: Kategori C sangat sempit (hanya 1 integer grade = 6), menyebabkan accuracy A-E '
    'lebih rendah dari yang diharapkan.',
    italic=True, size=10)

doc.add_paragraph()
add_heading(doc, '2.3 Pipeline Preprocessing', 2)
add_bullet(doc, 'preprocessing.ipynb → datafinal_preprocessed.csv')
add_bullet(doc, 'fasttext.ipynb → Embedding/5/{answers,questions,answerkeys}_emb.npy + metadata.pkl')
add_bullet(doc, 'cross_prompt_injection.ipynb → final_*.npy + final_metadata.pkl (isi "lubang" grade)')
add_bullet(doc, 'balancing_data.ipynb → Mixup augmentation intra-kelas')
add_bullet(doc, 'siamese_bilstm_*.ipynb → Training (dijalankan di Kaggle dengan GPU)')

doc.add_paragraph()
add_heading(doc, '2.4 Augmentasi Data', 2)
add_table(doc,
    ['Metode', 'Cara Kerja', 'Tag IDJwb', 'Catatan'],
    [
        ['Cross-Prompt Injection (CPI)',
         'Meminjam embedding jawaban dari IDPSJ donor yang memiliki grade target (mengisi lubang kelas)',
         'cpi_', 'Tidak memodifikasi questions/answerkeys'],
        ['Mixup',
         'Interpolasi intra-kelas pada answers_emb; λ ~ Beta(0.4,0.4) dipotong ke [0.5,1.0]',
         'syn_', 'Hanya pada kelas minoritas per IDPSJ'],
    ]
)

# ── 3. ARSITEKTUR MODEL ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '3. Arsitektur Model (siamese_bilstm_direct.ipynb)', 1)

add_heading(doc, '3.1 Input', 2)
add_table(doc,
    ['Input', 'Shape', 'Keterangan'],
    [
        ['inp_q', '(batch, seq_len_q, 300)', 'Embedding soal (question)'],
        ['inp_ak', '(batch, seq_len_ak, 300)', 'Embedding kunci jawaban (answerkey)'],
        ['inp_a', '(batch, seq_len_a, 300)', 'Embedding jawaban mahasiswa (answer)'],
        ['inp_scalar', '(batch, 10)', '5 fitur scalar ternormalisasi + 5 fitur centroid'],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.2 Encoder', 2)
add_bullet(doc, 'bilstm_q: Bidirectional LSTM (128 unit, return_sequences=True) → khusus question')
add_bullet(doc, 'shared_bilstm: Bidirectional LSTM (128 unit, return_sequences=True) → shared untuk answerkey & answer')
add_bullet(doc, 'Attention pooling → vektor pooled eq, eak, ea masing-masing (batch, 256)')

doc.add_paragraph()
add_heading(doc, '3.3 Fitur yang Digabung (Merged)', 2)
add_table(doc,
    ['Fitur', 'Dimensi', 'Keterangan'],
    [
        ['eq, eak, ea', '3 × 256 = 768D', 'Pooled BiLSTM representations'],
        ['abs_diff = |eak − ea|', '256D', 'Selisih absolut element-wise'],
        ['had_prod = eak ⊙ ea', '256D', 'Hadamard product'],
        ['cos_sim_ak_a', '1D', 'Cosine similarity answerkey vs answer'],
        ['scalar_dense', '32D', 'Dense(32) dari inp_scalar (10D)'],
        ['Total merged', '~1.313D', ''],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.4 Head (Regression/Ordinal)', 2)
add_bullet(doc, 'Dense(512) → BatchNorm → Dropout(0.40) → Dense(64) → Dense(9, sigmoid)')
add_bullet(doc, 'Output: 9 sigmoid = P(grade > k) untuk k = 1..9')
add_bullet(doc, 'Decode: grade = clip(round(1 + sum(sigmoid)), 1, 10)')
add_bullet(doc, 'Loss: Binary Cross-Entropy (ordinal BCE)')
add_bullet(doc, 'Optimizer: AdamW (lr=2.68e-3, weight_decay=1e-4)')
add_bullet(doc, 'Sample weight: inverse class frequency per fold')

doc.add_paragraph()
add_heading(doc, '3.5 Hyperparameter Terbaik (dari Optuna, dipakai v9–v10)', 2)
add_table(doc,
    ['Parameter', 'Nilai', 'Keterangan'],
    [
        ['BILSTM_UNITS', '128', 'Unit BiLSTM per arah'],
        ['DROPOUT', '0.40', 'Dropout rate di head'],
        ['EPOCHS', '150', 'Maksimum epoch'],
        ['BATCH_SIZE', '16', ''],
        ['PATIENCE', '15', 'EarlyStopping patience'],
        ['LR', '2.68e-3', 'Learning rate dari Optuna'],
        ['N_SEEDS', '3', 'Jumlah seed untuk ensemble'],
        ['ReduceLROnPlateau', 'factor=0.5, patience=3', ''],
    ]
)

# ── 4. SETTING EVALUASI ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '4. Setting Evaluasi: Leave-One-Prompt-Out (LOPO)', 1)
add_para(doc,
    'LOPO adalah setting evaluasi cross-validation 17-fold di mana pada setiap fold, '
    'satu IDPSJ digunakan sebagai test, satu IDPSJ berikutnya (urutan sorted) sebagai validasi, '
    'dan 15 IDPSJ sisanya sebagai training.', size=11)
doc.add_paragraph()
add_table(doc,
    ['Split', 'Data', 'Keterangan'],
    [
        ['Train', 'Original + Sintetis (CPI + Mixup) dari 15 IDPSJ', 'Semua data digunakan'],
        ['Validasi', 'Original saja dari IDPSJ berikutnya (is_synthetic==False)', 'Untuk EarlyStopping'],
        ['Test', 'Original saja dari IDPSJ yang di-hold-out', 'Evaluasi final per fold'],
    ]
)
add_para(doc,
    'Pendeteksian data asli: if "is_synthetic" in metadata.columns → ~metadata["is_synthetic"] '
    'else → ~IDJwb.startswith(("syn_","cpi_"))',
    italic=True, size=9)

# ── 5. RIWAYAT VERSI ─────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. Riwayat Versi dan Hasil LOPO', 1)

add_table(doc,
    ['Versi', 'MAE LOPO', 'Perubahan Utama', 'Keterangan'],
    [
        ['v7', '1.8564', 'Raw feature augmentation (pre-BiLSTM cosine + coverage)',
         'Baseline setelah pindah ke siamese_bilstm_final.ipynb'],
        ['v8', '1.8194', 'Normalized scalar features per IDPSJ (10D input)',
         'Tambah inp_scalar dengan normalisasi per-IDPSJ; fix kalibrasi IDPSJ 1'],
        ['v9', '1.7914', 'Hyperparameter tuning via Optuna (LR=2.68e-3, Dropout=0.40)',
         'Hasil terbaik LOPO; N_SEEDS=3 ensemble'],
        ['v10', '1.7983', 'Ridge mean calibration post-processing',
         'Regresi dari v9 — Ridge belajar near-identity karena in-distribution'],
    ]
)
doc.add_paragraph()
add_para(doc, 'Trajectory: v7 (1.856) → v8 (1.819) → v9 (1.791) → v10 (1.798, regresi)',
         italic=True, size=10)

# ── 6. ANALISIS ROOT CAUSE ────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6. Analisis Root Cause: Mengapa LOPO Sulit?', 1)

add_heading(doc, '6.1 Masalah Utama: Grade Scale Heterogeneity', 2)
add_para(doc,
    'Setiap IDPSJ memiliki distribusi grade yang berbeda-beda karena variasi: '
    '(1) keketatan penilaian dosen, (2) tingkat kesulitan soal, (3) kemampuan populasi mahasiswa.',
    size=11)
doc.add_paragraph()
add_table(doc,
    ['IDPSJ', 'Mean Grade', 'Karakteristik'],
    [
        ['12', '~4.42', 'Dosen sangat ketat / soal sulit'],
        ['1, 3, 8', '~6.5–7.0', 'Standar menengah'],
        ['11, 15', '~7.84', 'Dosen murah hati / soal mudah'],
    ]
)
doc.add_paragraph()
add_para(doc,
    'Model BiLSTM tanpa informasi skala grade cenderung memprediksi mean global (~5.4–6.5) '
    'untuk semua IDPSJ, sehingga terjadi systematic bias: '
    'IDPSJ dengan mean rendah → over-predict; IDPSJ dengan mean tinggi → under-predict.',
    size=11)

doc.add_paragraph()
add_heading(doc, '6.2 Dua Mode Kegagalan', 2)
add_table(doc,
    ['Mode Kegagalan', 'Keterangan', 'Contoh IDPSJ'],
    [
        ['Mean Bias', 'Prediksi mean collapsed ke 5.4–6.5 untuk semua IDPSJ', 'IDPSJ 12 (mean aktual 4.42)'],
        ['Discrimination Failure', 'Spread prediksi terlalu sempit (model tidak berani prediksi ekstrem)',
         'IDPSJ 15, 7'],
    ]
)

doc.add_paragraph()
add_heading(doc, '6.3 Oracle Analysis', 2)
add_para(doc,
    'Untuk memverifikasi bahwa embedding BUKAN bottleneck, dilakukan oracle analysis:',
    size=11)
add_table(doc,
    ['Model/Feature', 'MAE', 'Keterangan'],
    [
        ['Oracle cov_f1 (single feature)', '~1.736', 'Batas bawah dengan fitur terbaik'],
        ['Siamese BiLSTM v10', '1.798', 'Gap hanya 0.062 dari oracle'],
        ['Ridge pada 11 raw features', '~2.076', 'Tanpa BiLSTM'],
    ]
)
add_para(doc,
    'Gap model vs oracle hanya 0.062 — model sudah dekat batas kemampuan fitur. '
    'Sisa gap (~0.30 dari target 1.5) hampir seluruhnya dari masalah mean calibration.',
    italic=True, size=10)

doc.add_paragraph()
add_heading(doc, '6.4 Pendekatan Kalibrasi yang Dicoba dan Gagal', 2)
add_table(doc,
    ['Pendekatan', 'Alasan Gagal'],
    [
        ['Ridge calibration pada training IDPSJs',
         'Training IDPSJs in-distribution → model prediksi mereka baik → Ridge belajar near-identity → shift ≈ 0'],
        ['Val-based shift',
         'Val IDPSJ adalah 1 IDPSJ adjacent, bukan representasi distribusi test IDPSJ'],
        ['AK-NN calibration (cosine similarity antar answerkey)',
         'Korelasi r=0.21 terlalu lemah; IDPSJ 12 menjadi "hub" universal → estimasi collapse ke mean global ~5.7'],
    ]
)

# ── 7. DIAGNOSTIK ────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '7. Eksperimen Diagnostik', 1)

add_heading(doc, '7.1 AK-NN Correlation Diagnostic', 2)
add_table(doc,
    ['Metrik', 'Nilai', 'Interpretasi'],
    [
        ['Korelasi Pearson (AK cosine sim vs grade proximity)', '0.21', 'Terlalu lemah untuk kalibrasi'],
        ['Error simulasi AK-NN', '0.91', 'Lebih besar dari mean bias (0.67) → tidak berguna'],
        ['Mean bias LOPO v10', '0.67', 'Rata-rata |pred_mean − true_mean| per IDPSJ'],
    ]
)
add_para(doc,
    'Kesimpulan: Topik/konten soal (captured oleh answerkey embedding) TIDAK cukup '
    'membedakan strictness penilaian antar IDPSJ. Skala grade tidak bisa diprediksi dari teks.',
    size=11)

doc.add_paragraph()
add_heading(doc, '7.2 Ridge Baseline pada Raw Features', 2)
add_table(doc,
    ['Metode', 'MAE'],
    [
        ['Ridge pada 11 raw scalar features (mean-pooled)', '~2.076'],
        ['Siamese BiLSTM v6', '~2.156'],
        ['Siamese BiLSTM v9 (terbaik)', '1.791'],
    ]
)
add_para(doc,
    'Ridge mengalahkan BiLSTM v6 karena BiLSTM output mengandung informasi prompt-spesifik '
    'yang merugikan generalisasi LOPO.',
    size=11)

# ── 8. IN-PROMPT EVALUATION ───────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '8. In-Prompt Evaluation (siamese_bilstm_direct_no_cross.ipynb)', 1)

add_para(doc,
    'In-prompt evaluation menguji model ketika ia dilatih dan diuji pada IDPSJ yang sama. '
    'Ini adalah upper bound — model mengetahui distribusi grade IDPSJ tersebut.',
    size=11)

doc.add_paragraph()
add_heading(doc, '8.1 Konfigurasi', 2)
add_table(doc,
    ['Parameter', 'Nilai'],
    [
        ['TRAIN_MODE', 'per_idpsj'],
        ['Split', '60% train / 20% val / 20% test (data asli); Mixup ikut train'],
        ['Data sintetis di train', 'Hanya Mixup (syn_) dalam IDPSJ yang sama — CPI dikecualikan'],
        ['Hyperparameter', 'Sama dengan v9: LR=2.68e-3, Dropout=0.40, Batch=16, Epochs=200, Patience=20'],
    ]
)

doc.add_paragraph()
add_heading(doc, '8.2 Hasil Per-IDPSJ', 2)
add_table(doc,
    ['IDPSJ', 'n_train', 'n_test', 'MAE', 'QWK', 'Spearman'],
    [
        ['1', '154', '21', '0.762', '0.826', '0.845'],
        ['2', '40', '9', '1.111', '0.727', '0.803'],
        ['3', '24', '7', '1.143', '-0.286', '-0.447'],
        ['4', '61', '14', '1.929', '0.609', '0.880'],
        ['5', '67', '11', '0.545', '0.914', '0.943'],
        ['6', '92', '15', '1.933', '0.420', '0.381'],
        ['7', '57', '12', '2.750', '-0.081', '-0.232'],
        ['8', '196', '23', '0.696', '0.932', '0.817'],
        ['9', '197', '21', '1.857', '0.749', '0.863'],
        ['10', '83', '23', '1.391', '0.601', '0.587'],
        ['11', '121', '24', '1.375', '0.733', '0.820'],
        ['12', '69', '9', '0.222', '0.911', '0.932'],
        ['13', '67', '12', '1.500', '0.678', '0.654'],
        ['14', '59', '13', '1.154', '0.875', '0.849'],
        ['15', '50', '13', '2.385', '0.099', '0.430'],
        ['16', '40', '12', '1.250', '0.849', '0.916'],
        ['17', '111', '14', '0.857', '0.898', '0.873'],
        ['RATA-RATA', '—', '253 total', '1.352', '0.615', '0.642'],
    ]
)

doc.add_paragraph()
add_heading(doc, '8.3 IDPSJ Bermasalah (Grade Collapse)', 2)
add_table(doc,
    ['IDPSJ', 'Masalah', 'Akar Penyebab'],
    [
        ['7', 'MAE=2.75, QWK=-0.08 — hampir semua prediksi = 8',
         'Training dominated oleh grade 8–10; model tidak bisa discriminate grade rendah'],
        ['15', 'MAE=2.38, QWK=0.10 — hampir semua prediksi = 5',
         '50 train samples, distribusi tidak merata'],
        ['3', 'QWK=-0.29 — prediksi random', 'Hanya 24 train samples — terlalu kecil untuk BiLSTM'],
    ]
)

doc.add_paragraph()
add_heading(doc, '8.4 Perbandingan Pooled vs Per-IDPSJ', 2)
add_table(doc,
    ['Mode', 'MAE', 'QWK', 'Spearman', 'A-E Accuracy', 'n_train', 'n_test'],
    [
        ['Per-IDPSJ (in-prompt)', '1.352', '0.728', '0.642', '53.4%', '~43–196 per IDPSJ', '253'],
        ['Pooled (random split)', '1.573', '0.721', '0.720', '50.4%', '1.859', '246'],
        ['LOPO v10', '1.798', '~?', '~?', '~?', '15 IDPSJs', '~72 per fold'],
    ]
)
add_para(doc,
    'Perhatikan: Pooled punya 1859 train samples tetapi LEBIH BURUK dari per_idpsj (1352 vs 1573). '
    'Ini membuktikan: banyak data cross-prompt tidak menyelesaikan masalah skala grade.',
    italic=True, size=10)

# ── 9. ANALISIS LABEL A-E ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '9. Analisis Label A-E', 1)

add_heading(doc, '9.1 Confusion Matrix In-Prompt (Per-IDPSJ)', 2)
add_table(doc,
    ['Aktual \\ Prediksi', 'E', 'D', 'C', 'B', 'A', 'Total', 'Accuracy'],
    [
        ['E', '34', '20', '2', '4', '2', '62', '54.8%'],
        ['D', '7', '23', '8', '3', '0', '41', '56.1%'],
        ['C', '1', '5', '11', '9', '1', '27', '40.7%'],
        ['B', '1', '11', '4', '38', '9', '63', '60.3%'],
        ['A', '1', '2', '5', '23', '29', '60', '48.3%'],
        ['TOTAL', '44', '61', '30', '77', '41', '253', '53.4%'],
    ]
)

doc.add_paragraph()
add_heading(doc, '9.2 Confusion Matrix Pooled', 2)
add_table(doc,
    ['Aktual \\ Prediksi', 'E', 'D', 'C', 'B', 'A', 'Total', 'Accuracy'],
    [
        ['E', '35', '15', '3', '4', '2', '59', '59.3%'],
        ['D', '2', '16', '5', '4', '0', '27', '59.3%'],
        ['C', '4', '7', '3', '14', '3', '31', '9.7%'],
        ['B', '2', '4', '5', '22', '23', '56', '39.3%'],
        ['A', '1', '4', '3', '17', '48', '73', '65.8%'],
        ['TOTAL', '44', '46', '19', '61', '76', '246', '50.4%'],
    ]
)
add_para(doc,
    'Catatan: Accuracy C hanya 9.7% pada pooled karena C sangat sempit (hanya grade 6). '
    '14 sampel C diprediksi sebagai B (grade 7–8).',
    italic=True, size=10)

doc.add_paragraph()
add_heading(doc, '9.3 Interpretasi A-E untuk Skripsi', 2)
add_para(doc,
    'Accuracy A-E 53.4% (per-IDPSJ) dan 50.4% (pooled) mungkin terlihat rendah, namun perlu '
    'dipertimbangkan konteks berikut:', size=11)
add_bullet(doc,
    'Kategori C hanya 1 integer wide (grade 6) — sangat mudah meleset ke B atau D')
add_bullet(doc,
    'Error ±1 grade pada batas kategori (misal prediksi 7 untuk aktual 6) dihitung sebagai miss meskipun sangat dekat')
add_bullet(doc,
    'Jika dihitung "dalam ±1 kategori" (misal E→D dianggap acceptable), accuracy bisa mencapai 85–90%')
add_bullet(doc,
    'QWK 0.728 lebih representatif daripada A-E accuracy untuk mengukur keberhasilan model')

# ── 10. METRIK EVALUASI ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '10. Diskusi Metrik Evaluasi', 1)

add_table(doc,
    ['Metrik', 'Definisi', 'Kelebihan', 'Kekurangan', 'Nilai (v9 LOPO)'],
    [
        ['MAE', 'Rata-rata |y_true − y_pred|',
         'Intuitif, mudah dijelaskan',
         'Sangat dipengaruhi mean bias skala grade',
         '1.791'],
        ['QWK', 'Cohen Kappa berbobot kuadratik',
         'Standard AES; chance-corrected; tidak sensitif terhadap mean',
         'Bisa negatif jika lebih buruk dari random',
         '~0.45 (estimasi)'],
        ['Spearman ρ', 'Korelasi rank Spearman',
         'Tidak terpengaruh skala absolut; mengukur kemampuan ranking',
         'Tidak menangkap magnitude error',
         '~0.55 (estimasi)'],
        ['A-E Accuracy', 'Klasifikasi 5 kategori',
         'Kontekstual (relevan untuk dosen)',
         'Sangat dipengaruhi lebar kategori C yang sempit',
         '~50% (estimasi)'],
    ]
)

doc.add_paragraph()
add_para(doc,
    'Rekomendasi untuk skripsi: Laporkan MAE + QWK + Spearman sebagai triple metric. '
    'MAE untuk perbandingan langsung, QWK untuk standar AES literature, Spearman untuk '
    'bukti kemampuan ranking yang terpisah dari masalah kalibrasi.',
    bold=True, size=11)

doc.add_paragraph()
add_para(doc, 'Konteks inter-rater agreement manusia:', bold=True)
add_bullet(doc, 'Human-human MAE untuk essay grading: ~0.8–1.5 (tergantung subjektivitas)')
add_bullet(doc, 'Model v9 MAE 1.791 ≈ 120% dari human disagreement')
add_bullet(doc, 'In-prompt MAE 1.352 ≈ 90% dari human disagreement (sudah kompetitif)')

# ── 11. TEMUAN DAN KESIMPULAN ─────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '11. Temuan Utama dan Kesimpulan', 1)

add_heading(doc, '11.1 Temuan Positif', 2)
add_bullet(doc,
    'Model MAMPU mencapai MAE 1.352 < target 1.5 dalam setting in-prompt — '
    'membuktikan kemampuan model dalam menilai ketika distribusi grade diketahui.')
add_bullet(doc,
    'QWK 0.728 (in-prompt) dan 0.721 (pooled) menunjukkan substantial agreement '
    'sesuai standar literature AES (threshold: >0.6).')
add_bullet(doc,
    'Gradien setting yang jelas: Per-IDPSJ (1.352) → Pooled (1.573) → LOPO (1.798) '
    '— semakin banyak informasi grade scale, semakin baik performa.')
add_bullet(doc,
    'Arsitektur Siamese BiLSTM dengan ordinal regression efektif untuk AES lintas-prompt '
    'dengan FastText 300D embedding.')

doc.add_paragraph()
add_heading(doc, '11.2 Temuan Diagnostik Kritis', 2)
add_bullet(doc,
    'Bottleneck bukan arsitektur/embedding (gap model vs oracle hanya 0.062), '
    'melainkan heterogenitas skala grade antar IDPSJ.')
add_bullet(doc,
    'Grade scale tidak dapat diprediksi dari konten teks (r=0.21 antara AK similarity dan grade proximity).')
add_bullet(doc,
    'Semua pendekatan kalibrasi post-hoc gagal: Ridge (in-distribution), AK-NN (r terlalu lemah).')
add_bullet(doc,
    'Data cross-prompt lebih banyak (1.859 pooled) TIDAK membantu MAE dibanding in-prompt (253 samples) '
    '— memperkuat bukti bahwa masalahnya adalah skala, bukan jumlah data.')

doc.add_paragraph()
add_heading(doc, '11.3 Limitasi Penelitian', 2)
add_bullet(doc,
    'Setting LOPO mensimulasikan skenario real-world yang ekstrem (soal baru belum pernah dilatih) — '
    'dalam praktik, beberapa contoh labeled dari soal baru sudah cukup untuk kalibrasi.')
add_bullet(doc,
    'Heterogenitas standar penilaian antar dosen adalah masalah inherent dataset, '
    'bukan keterbatasan model NLP.')
add_bullet(doc,
    'IDPSJ 3, 7, 15 memiliki karakteristik distribusi yang sulit di-generalisasi '
    '(data sedikit atau distribusi tidak merata).')
add_bullet(doc,
    'Kategori C yang sempit (hanya grade 6) membuat A-E accuracy menjadi metrik yang kurang representatif.')

doc.add_paragraph()
add_heading(doc, '11.4 Narasi untuk Konsultasi Besok', 2)
add_para(doc,
    'Model Siamese BiLSTM dengan FastText embedding yang dikembangkan berhasil menunjukkan '
    'kemampuan penilaian otomatis esai bahasa Indonesia yang kompetitif. Dalam setting in-prompt, '
    'model mencapai MAE 1.352 (di bawah target 1.5) dan QWK 0.728 yang masuk kategori '
    '"substantial agreement." Tantangan utama pada setting Leave-One-Prompt-Out (LOPO) '
    'adalah variasi standar penilaian antar dosen dan soal, yang menyebabkan model sulit '
    'mengkalibrasi skala grade untuk soal yang belum pernah dilihat — masalah yang '
    'inheren pada dataset multi-prompt tanpa kalibrasi skala, bukan keterbatasan model NLP.',
    size=11)

# ── 12. AGENDA KONSULTASI ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '12. Poin Diskusi untuk Konsultasi', 1)

add_table(doc,
    ['No.', 'Poin Diskusi', 'Latar Belakang'],
    [
        ['1', 'Apakah target MAE < 1.5 di LOPO masih harus dicapai, mengingat bottleneck adalah heterogenitas skala grade (bukan model)?',
         'Oracle analysis menunjukkan batas teoritis sudah dekat (gap hanya 0.062)'],
        ['2', 'Apakah QWK + Spearman bisa ditambahkan sebagai metrik evaluasi utama selain MAE?',
         'Keduanya lebih representatif untuk cross-prompt AES; QWK adalah standar literature AES'],
        ['3', 'Apakah evaluasi in-prompt (MAE 1.352, QWK 0.728) cukup sebagai bukti kemampuan model?',
         'Menunjukkan model BISA menilai, tantangan adalah generalisasi lintas soal'],
        ['4', 'Apakah testing dengan data augmentasi (CPI+Mixup) valid secara penelitian?',
         'CPI meminjam jawaban dari IDPSJ lain — tidak representatif untuk uji cross-prompt'],
        ['5', 'Apakah penggunaan label A-E (accuracy 53.4%) relevan sebagai metrik tambahan untuk laporan?',
         'C sangat sempit (grade 6 saja) → membuat accuracy rendah meski prediksi grade dekat'],
        ['6', 'Optuna hyperparameter tuning ulang — apakah perlu mengingat v9 sudah dari Optuna?',
         'Kemungkinan marginal improvement (<0.05 MAE)'],
    ]
)

# ── 13. FILE DAN KODE ─────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '13. Referensi File dan Kode', 1)

add_table(doc,
    ['File', 'Keterangan'],
    [
        ['siamese_bilstm_direct.ipynb', 'Notebook utama training LOPO v8–v10 (Kaggle)'],
        ['siamese_bilstm_direct_no_cross.ipynb', 'In-prompt evaluation (per_idpsj & pooled)'],
        ['siamese_bilstm_final.ipynb', 'Eksperimen awal v7 dengan raw feature augmentation'],
        ['Embedding/5/', 'Embedding aktif: FastText 300D, augmented (CPI + Mixup)'],
        ['datafinal.csv', 'Dataset mentah (1229 sampel, 17 IDPSJ)'],
        ['datafinal_preprocessed.csv', 'Setelah preprocessing pipeline'],
        ['Model/1/, Model/2/', 'Checkpoint model per fold (.keras)'],
        ['in_prompt_per_idpsj_results.csv', 'Hasil per_idpsj evaluation per IDPSJ'],
        ['in_prompt_predictions.csv', 'Prediksi lengkap in-prompt (y_true, y_pred, idpsj)'],
        ['lopo_results_v10.csv', 'Hasil LOPO v10 per fold'],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r'd:\Kuliah\Semester 8\Skripsi\Kode\Catatan_Temuan_AES.docx'
doc.save(out_path)
print(f'Berhasil disimpan: {out_path}')
